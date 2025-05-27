from docx import Document

# ─────────────────────────────────────────────────────────────
#  ORIGINAL FUNCTION — Base DOCX handler
# ─────────────────────────────────────────────────────────────

def extract_text_from_docx(file_path):
    """
    Extracts meaningful text from a .docx file, including:
    - Paragraphs
    - Tables
    - (Optionally) metadata (if needed in the future)

    Args:
        file_path (str): Path to the .docx file.

    Returns:
        str: Clean, readable, and structured text.
    """
    try:
        doc = Document(file_path)
        extracted = []

        # --- Extract plain text paragraphs ---
        extracted.append("### Document Body ###\n")
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted.append(text)

        # --- Extract text from tables ---
        if doc.tables:
            extracted.append("\n### Tables ###\n")
            for i, table in enumerate(doc.tables):
                extracted.append(f"[Table {i+1}]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    extracted.append("\t".join(row_data))
                extracted.append("")  # Newline between tables

        # --- Placeholder: Metadata support ---
        # props = doc.core_properties
        # extracted.append(f"Author: {props.author}")
        # extracted.append(f"Title: {props.title}")
        # extracted.append("")

        return "\n".join(extracted).strip()

    except Exception as e:
        print(f"[ERROR] Failed to extract DOCX: {file_path} – {e}")
        return ""

# ─────────────────────────────────────────────────────────────
#  BONUS FUNCTION — Enhanced DOCX handler with metadata
# ─────────────────────────────────────────────────────────────

import logging

# Logger setup
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')

def extract_text_from_docx_enhanced(file_path, include_metadata=True):
    """
    Enhanced version of .docx text extraction:
    - Extracts paragraphs and tables
    - Optionally includes core document metadata
    - Uses structured logging

    Args:
        file_path (str): Path to the .docx file.
        include_metadata (bool): If True, include metadata like author/title.

    Returns:
        str: Structured document content with optional metadata.
    """
    try:
        doc = Document(file_path)
        extracted = []

        # Optional metadata block
        if include_metadata:
            props = doc.core_properties
            extracted.append("### Document Metadata ###")
            extracted.append(f"Title: {props.title or 'N/A'}")
            extracted.append(f"Author: {props.author or 'N/A'}")
            extracted.append(f"Created: {props.created.strftime('%Y-%m-%d') if props.created else 'N/A'}")
            extracted.append("")

        # Extract body text
        extracted.append("### Document Body ###")
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted.append(text)

        # Extract tables, if any
        if doc.tables:
            extracted.append("\n### Tables ###")
            for i, table in enumerate(doc.tables):
                extracted.append(f"[Table {i+1}]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        extracted.append("\t".join(row_data))
                extracted.append("")

        return "\n".join(extracted).strip()

    except Exception as e:
        logger.error(f"Enhanced DOCX extraction failed for {file_path}: {e}")
        return ""
